import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

import re
import io
import datetime
import pytz
from loguru import logger

from ..re_pattern import patterns

tz = pytz.timezone('Asia/Shanghai')

def shuttle_text(shuttle):
    full_text = ''
    last_end = 0
    for i in shuttle:
        full_text += i.text
        last_end = last_end+len(i.text)
    return full_text

def paragraph_process(p):
    begin = 0
    inner_content = list()
    for item in p.iter_inner_content():
        if isinstance(item, docx.text.hyperlink.Hyperlink):
            for i in item.runs:
                inner_content.append(i)
        else:
            inner_content.append(item)
    for end in range(len(inner_content)):
        for pa in patterns:
            full_text = shuttle_text(inner_content[begin:end+1])
            r = re.search(pa, full_text)
            if r:
                index = r.start()
                while index >= len(inner_content[begin].text):
                    index -= len(inner_content[begin].text)
                    begin += 1
                run = inner_content[begin:end+1]
                run_c = re.search(pa, run[0].text)
                if run_c:
                    if len(run_c.group()) <= 8:
                        run[0].text = 'X' * len(run_c.group())
                    else:
                        run[0].text = run[0].text[:run_c.start()+3] + 'X'*(len(run_c.group())-6) + run[0].text[run_c.end()-3:]
                else:
                    new_text = shuttle_text(run)
                    run_p = re.search(pa, new_text)
                    if run_p:
                        if len(run_p.group()) <= 8:
                            new_text = 'X' * len(run_p.group())
                        else:
                            new_text = new_text[:run_p.start()+3] + 'X' * (len(run_p.group())-6) + new_text[run_p.end()-3:]
                        index = run_p.start()
                        run[0].text = new_text[:len(run[0].text)]
                        index = len(run[0].text)

                        for i in run[1:-1]:
                            i.text=new_text[index:index+len(i.text)]
                            index += len(i.text)
                        
                        run[-1].text = new_text[index:]
                    else:
                        continue
                begin = end
    return p

def table_process(t):
    for c in t.columns:
        for cell in c.cells:
            for p in cell.paragraphs:
                p = paragraph_process(p)
    return t

def hyperlink_process(file):
    rels = file.part.rels
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            for p in patterns:
                if re.search(p,rels[rel]._target):
                     rels[rel]._target = '[原始敏感链接已清除] [ORIGIN SENSITIVE HYPERLINK REMOVED]'

def docx_process(file_t, file_id, user, org):
    file = docx.Document(file_t)
    file_t.close()
    
    for p in file.paragraphs:
        p = paragraph_process(p)

    for t in file.tables:
        for c in t.columns:
            for cell in c.cells:
                for p in cell.paragraphs:
                    p = paragraph_process(p)

    for s in file.sections:
        items = list()
        for i in s.header.iter_inner_content():
            items.append(i)
        for i in s.footer.iter_inner_content():
            items.append(i)
        for i in items:
            if isinstance(i, docx.text.paragraph.Paragraph):
                i = paragraph_process(i)
            else:
                i = table_process(i)
    
    hyperlink_process(file)

    time = datetime.datetime.now() - datetime.timedelta(hours=8)
    file.core_properties.author = "{}@{};{}@{}".format(user.username, org.name, user.uid, org.oid)
    file.core_properties.last_modified_by = "Sensitive Data Recognizer@Xiamen_Torch_HiTech_IDZ"
    file.core_properties.created = time
    file.core_properties.modified = time
    file.core_properties.last_printed = time
    file.core_properties.revision = 25565
    file.core_properties.title = file_id
    file.core_properties.comments = "由文件脱敏系统处理部分关键文字信息。\n@Sensitive Data Recognizer\n@Xiamen_Torch_HiTech_IDZ\n@Chengyi University College, Jimei University"
    new_file = io.BytesIO()
    file.save(new_file)
    new_file.seek(0)
    return new_file
    
